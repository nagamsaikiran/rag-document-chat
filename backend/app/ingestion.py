"""Document loading and chunking: PDF, DOCX, TXT, MD, HTML.

Design note (see README): we use a *recursive* splitter that prefers to break
on paragraph -> line -> sentence -> word boundaries before resorting to a hard
character cut. Naive fixed-size slicing splits mid-sentence and hurts retrieval
quality; this keeps chunks semantically coherent while bounding their size.

PDF extras:
  - optional vision ingestion (page image -> markdown, reads tables/figures)
  - repeated header/footer lines are stripped before chunking so page
    furniture ("Confidential", page numbers, running titles) doesn't pollute
    retrieval.
"""
from collections import Counter
from dataclasses import dataclass, field
from html.parser import HTMLParser
from typing import List

from pypdf import PdfReader

from app.config import get_settings

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".html", ".htm")


@dataclass
class Chunk:
    text: str
    source: str  # filename
    page: int    # 1-indexed page number (1 for unpaged formats)
    chunk_id: str = field(default="")


# Instruction given to the vision model for each page image.
_VISION_PROMPT = (
    "Transcribe this document page into clean Markdown. Rules:\n"
    "- Preserve reading order.\n"
    "- Render any tables as proper Markdown tables, keeping rows and columns.\n"
    "- For charts, graphs, photos, logos or diagrams, insert a concise description "
    "in the form: [Figure: <what it shows, including any numbers/labels>].\n"
    "- Do not add commentary or invent content. Output only the page content."
)


def _check_page_limit(n_pages: int, vision: bool = False) -> None:
    settings = get_settings()
    limit = settings.max_pages
    if vision:
        # Vision costs one model call per page; keep a tighter cap.
        limit = min(limit, settings.max_pages_vision)
    if n_pages > limit:
        raise ValueError(
            f"PDF has {n_pages} pages; the limit is {limit}"
            f"{' in vision mode' if vision else ''}. "
            "Split the document or raise MAX_PAGES."
        )


# --------------------------------------------------------------- PDF loading

def _strip_repeated_lines(pages: List[tuple[int, str]]) -> List[tuple[int, str]]:
    """Remove header/footer lines that repeat across most pages.

    A line counts as page furniture when it appears in the first or last two
    lines of >= 60% of pages (min 3 pages). Purely numeric lines (page numbers)
    are also dropped.
    """
    if len(pages) < 3:
        return pages
    edge_counts: Counter[str] = Counter()
    for _, text in pages:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        for ln in set(lines[:2] + lines[-2:]):
            edge_counts[ln] += 1
    threshold = max(2, int(len(pages) * 0.6))
    furniture = {ln for ln, n in edge_counts.items() if n >= threshold}

    cleaned = []
    for no, text in pages:
        kept = []
        lines = [ln for ln in text.splitlines()]
        for idx, ln in enumerate(lines):
            s = ln.strip()
            at_edge = idx < 2 or idx >= len(lines) - 2
            if at_edge and s and (s in furniture or s.isdigit()):
                continue
            kept.append(ln)
        body = "\n".join(kept).strip()
        if body:
            cleaned.append((no, body))
    return cleaned or pages


def load_pdf_pages_text(path: str) -> List[tuple[int, str]]:
    """Fast path: extract the embedded text layer with pypdf."""
    reader = PdfReader(path)
    _check_page_limit(len(reader.pages))
    pages = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((i + 1, text))
    return _strip_repeated_lines(pages)


def load_pdf_pages_vision(path: str) -> List[tuple[int, str]]:
    """Multimodal path: render each page to an image and have a vision model
    transcribe it into Markdown — captures tables, charts, figures, and scanned
    text that the plain text layer misses."""
    import fitz  # PyMuPDF, imported lazily so text-only mode needs no dep

    from app.llm.factory import get_llm

    llm = get_llm()
    dpi = get_settings().multimodal_dpi
    pages = []
    doc = fitz.open(path)
    try:
        _check_page_limit(doc.page_count, vision=True)
        for i in range(len(doc)):
            png = doc[i].get_pixmap(dpi=dpi).tobytes("png")
            text = llm.transcribe_image(png, "image/png", _VISION_PROMPT).strip()
            if text:
                pages.append((i + 1, text))
    finally:
        doc.close()
    return pages


def load_pdf_pages(
    path: str, source_name: str, multimodal: bool | None = None
) -> List[tuple[int, str]]:
    """Return [(page_number, page_text)] for a PDF, skipping empty pages.

    Vision mode (tables/charts/figures/scanned) is used when enabled, otherwise
    the fast pypdf text path. `multimodal` overrides the global default per call
    (so the UI can toggle it per upload). If the provider has no vision support
    we cleanly fall back to text; real vision errors propagate so they surface.
    """
    use_vision = get_settings().multimodal if multimodal is None else multimodal
    if use_vision:
        try:
            return load_pdf_pages_vision(path)
        except NotImplementedError:
            pass  # provider has no vision; use the text path instead
    return load_pdf_pages_text(path)


# ------------------------------------------------------- other file formats

class _HTMLTextExtractor(HTMLParser):
    _SKIP = {"script", "style", "noscript", "head"}

    def __init__(self) -> None:
        super().__init__()
        self._parts: List[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip_depth += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data):
        if not self._skip_depth and data.strip():
            self._parts.append(data.strip())

    def text(self) -> str:
        return "\n".join(self._parts)


def _read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def load_document_pages(
    path: str, source_name: str, multimodal: bool | None = None
) -> List[tuple[int, str]]:
    """Dispatch by extension. Unpaged formats report everything as page 1."""
    name = source_name.lower()
    if name.endswith(".pdf"):
        return load_pdf_pages(path, source_name, multimodal)
    if name.endswith(".docx"):
        import docx  # python-docx, lazy import

        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n\n".join(parts).strip()
        return [(1, text)] if text else []
    if name.endswith((".html", ".htm")):
        extractor = _HTMLTextExtractor()
        extractor.feed(_read_text_file(path))
        text = extractor.text().strip()
        return [(1, text)] if text else []
    if name.endswith((".txt", ".md")):
        text = _read_text_file(path).strip()
        return [(1, text)] if text else []
    raise ValueError(
        f"Unsupported file type for '{source_name}'. "
        f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
    )


# ------------------------------------------------------------------ chunking

_SEPARATORS = ["\n\n", "\n", ". ", " "]


def _merge(pieces: List[str], sep: str, size: int, overlap: int) -> List[str]:
    """Greedily pack pieces into chunks <= size, carrying overlap between them."""
    chunks: List[str] = []
    current = ""
    for piece in pieces:
        if not piece:
            continue
        candidate = piece if not current else current + sep + piece
        if len(candidate) <= size:
            current = candidate
            continue
        if current:
            chunks.append(current)
        if overlap and chunks:
            tail = chunks[-1][-overlap:]
            current = tail + sep + piece
            if len(current) > size:  # overlap didn't fit; drop it
                current = piece
        else:
            current = piece
    if current:
        chunks.append(current)
    return chunks


def _recursive_split(
    text: str, size: int, overlap: int, _seps: List[str] | None = None
) -> List[str]:
    """True recursive split: paragraph -> line -> sentence -> word -> hard cut.

    Oversized fragments are split again with the *next* separator (this is the
    fix over a single-pass splitter, which hard-cuts any paragraph that has no
    top-level separator inside it).
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= size:
        return [text]
    seps = _SEPARATORS if _seps is None else _seps
    if not seps:
        # No separators left: hard slice with overlap.
        step = max(1, size - overlap)
        return [text[i:i + size] for i in range(0, len(text), step)]
    sep, rest = seps[0], seps[1:]
    if sep not in text:
        return _recursive_split(text, size, overlap, rest)
    pieces: List[str] = []
    for part in text.split(sep):
        if len(part) > size:
            pieces.extend(_recursive_split(part, size, overlap, rest))
        else:
            pieces.append(part)
    merged = _merge(pieces, sep, size, overlap)
    return [c.strip() for c in merged if c.strip()]


def chunk_file(
    path: str, source_name: str, multimodal: bool | None = None
) -> List[Chunk]:
    settings = get_settings()
    chunks: List[Chunk] = []
    for page_no, page_text in load_document_pages(path, source_name, multimodal):
        for j, piece in enumerate(
            _recursive_split(page_text, settings.chunk_size, settings.chunk_overlap)
        ):
            chunks.append(
                Chunk(
                    text=piece,
                    source=source_name,
                    page=page_no,
                    chunk_id=f"{source_name}::p{page_no}::c{j}",
                )
            )
    return chunks


def chunk_pdf(
    path: str, source_name: str, multimodal: bool | None = None
) -> List[Chunk]:
    """Back-compat alias (PDF was the only format in v1)."""
    settings = get_settings()
    chunks: List[Chunk] = []
    for page_no, page_text in load_pdf_pages(path, source_name, multimodal):
        for j, piece in enumerate(
            _recursive_split(page_text, settings.chunk_size, settings.chunk_overlap)
        ):
            chunks.append(
                Chunk(
                    text=piece,
                    source=source_name,
                    page=page_no,
                    chunk_id=f"{source_name}::p{page_no}::c{j}",
                )
            )
    return chunks
